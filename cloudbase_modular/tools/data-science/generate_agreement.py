from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.2)
    section.left_margin   = Cm(2.8)
    section.right_margin  = Cm(2.8)
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)

# Colour palette
NAVY  = RGBColor(0x0d, 0x1b, 0x3e)
GOLD  = RGBColor(0xB8, 0x90, 0x2E)
MID   = RGBColor(0x37, 0x41, 0x51)
DIM   = RGBColor(0x6b, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLUE_TINT = 'EDF6FF'
GOLD_TINT = 'FDF6E3'

# ── XML helpers ───────────────────────────────────────────────────────────────
def shade_paragraph(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

def shade_cell(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color='DDE3EA', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), sz)
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        borders.append(tag)
    tcPr.append(borders)

def add_rule(color='0d1b3e', sz=12, space_before=0, space_after=0):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), str(sz))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    return p

def spacer(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(pts)
    return p

# ── Text helpers ──────────────────────────────────────────────────────────────
def run(para, text, bold=False, italic=False, size=9.5,
        color=MID, font='Calibri'):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r

def body(text, indent=0.5, size=9.5, before=3, after=3):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(indent)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    run(p, text, size=size)
    return p

def section_heading(number, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(4)
    shade_paragraph(p, '0d1b3e')
    if number:
        run(p, f'  {number}.  ', bold=True, size=8, color=GOLD, font='Courier New')
    run(p, f'  {title.upper()}', bold=True, size=10, color=WHITE)
    return p

def sub(number, label, text, label_color=NAVY):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before      = Pt(4)
    p.paragraph_format.space_after       = Pt(4)
    if number:
        run(p, f'{number}  ', bold=True, size=8, color=GOLD, font='Courier New')
    if label:
        run(p, label, bold=True, size=9.5, color=label_color)
    run(p, text, size=9.5)
    return p

def bullet(text, indent=1.8):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Cm(indent)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)
    run(p, text, size=9.5)
    return p

def styled_table(headers, rows, col_widths=None):
    n = len(headers)
    tbl = doc.add_table(rows=1 + len(rows), cols=n)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'
    # Header
    hcells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hcells[i], '0d1b3e')
        p = hcells[i].paragraphs[0]
        run(p, h, bold=True, size=7.5, color=WHITE, font='Courier New')
        hcells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Rows
    for ri, row_data in enumerate(rows):
        fill = 'F8F9FB' if ri % 2 == 0 else 'FFFFFF'
        rcells = tbl.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            shade_cell(rcells[ci], fill)
            p = rcells[ci].paragraphs[0]
            bold_cell = (val == val.upper() and len(val) > 3 and ci == 0)
            run(p, str(val), bold=bold_cell, size=9, color=NAVY if bold_cell else MID)
            set_cell_borders(rcells[ci])
    # Widths
    if col_widths:
        for row in tbl.rows:
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = Inches(col_widths[ci])
    spacer(6)
    return tbl

# =============================================================================
# LETTERHEAD
# =============================================================================
p_brand = doc.add_paragraph()
p_brand.paragraph_format.space_before = Pt(0)
p_brand.paragraph_format.space_after  = Pt(2)
run(p_brand, 'ADHERENCE ', bold=False, size=22, color=NAVY, font='Georgia')
run(p_brand, 'Cartography', italic=True, size=22, color=GOLD, font='Georgia')

p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after  = Pt(4)
run(p_sub, 'ATLAS Clinical Platform   |   atlas.adherence.cc   |   info@adherence.cc',
    size=7.5, color=DIM, font='Courier New')

add_rule('B8902E', 18, 0, 0)
add_rule('0d1b3e', 6,  0, 0)
spacer(4)

# Document title
p_eye = doc.add_paragraph()
p_eye.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_eye.paragraph_format.space_before = Pt(10)
p_eye.paragraph_format.space_after  = Pt(6)
run(p_eye, 'ATLAS CLINICAL NETWORK SERVICES AGREEMENT',
    bold=True, size=8, color=GOLD, font='Courier New')

p_ttl = doc.add_paragraph()
p_ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ttl.paragraph_format.space_after = Pt(14)
run(p_ttl, 'Al Thiqa Pharmacy Network  |  United Arab Emirates',
    size=16, color=NAVY, font='Georgia')

add_rule('DDE3EA', 6, 0, 0)
spacer(4)

# Parties introduction
p_intro = doc.add_paragraph()
p_intro.paragraph_format.space_before = Pt(6)
p_intro.paragraph_format.space_after  = Pt(8)
run(p_intro, 'This Clinical Network Services Agreement (the "Agreement") is entered into as of ')
run(p_intro, '[DATE]', bold=True, color=GOLD)
run(p_intro, ' (the "Effective Date") by and between:')

def party_block(name, details):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run(p, name, bold=True, size=10, color=NAVY)
    run(p, details)
    return p

party_block('Adherence Cartography, Inc.',
    ', a company incorporated under the laws of California, United States, '
    'with its principal place of business at 100 Oceangate, 12th Floor, '
    'Long Beach, CA 90802 ("Adherence Inc."); and')

party_block('Al Thiqa Pharmacy Network',
    ', a company incorporated under the laws of the United Arab Emirates, '
    'with its principal place of business at [ADDRESS], Abu Dhabi, UAE ("Al Thiqa").')

spacer(6)
add_rule('DDE3EA', 4, 0, 0)

# =============================================================================
# RECITALS
# =============================================================================
section_heading('', 'RECITALS')

recitals = [
    ('WHEREAS,',
     'Adherence Inc. develops and operates ATLAS (Adherence Technology and Longitudinal '
     'Assessment System), a validated clinical platform for measuring and monitoring '
     'medication adherence using the Morisky Medication Adherence Scale (MMAS-8) and the '
     'Multidimensional Adherence Parameters (MAP) instrument;'),
    ('WHEREAS,',
     'Al Thiqa operates nine pharmacy locations across the United Arab Emirates, comprising '
     'one specialty pharmacy and eight retail pharmacy sites, and wishes to implement '
     'structured medication adherence assessment and monitoring across its network as part '
     'of its clinical pharmacy services programme;'),
    ('WHEREAS,',
     'the Parties recognise that the clinical dataset generated through this Agreement '
     'constitutes a significant research opportunity relating to medication adherence in '
     'the Gulf Arab population, and wish to formalise a research collaboration framework '
     'alongside the platform services described herein;'),
    ('NOW THEREFORE,',
     'in consideration of the mutual covenants and obligations set out herein, '
     'the Parties agree as follows:'),
]

for label, text in recitals:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(5)
    shade_paragraph(p, GOLD_TINT)
    run(p, label + '  ', bold=True, italic=True, color=NAVY)
    run(p, text, italic=True)

# =============================================================================
# 1. DEFINITIONS
# =============================================================================
section_heading(1, 'DEFINITIONS')
sub('1.1', '', 'In this Agreement the following terms shall have the meanings set out below:')

definitions = [
    ('"ATLAS Platform"',
     'means the web-based software platform operated by Adherence Inc. at atlas.adherence.cc, '
     'including all instruments, dashboards, reporting tools, data export functions, and associated infrastructure.'),
    ('"Workspace"',
     'means a siloed data environment within the ATLAS Platform assigned to a single Al Thiqa location, '
     'accessible only to authorised personnel at that location.'),
    ('"Network"',
     'means the collective of nine Al Thiqa pharmacy locations covered by this Agreement, as listed in Schedule A.'),
    ('"Instruments"',
     'means the validated psychometric tools licensed for use under this Agreement: MMAS-8 '
     '(Morisky Medication Adherence Scale, 8-item version) and MAP (Multidimensional Adherence Parameters).'),
    ('"Clinical Data"',
     'means de-identified or pseudonymised adherence scores, assessment records, and derived analytics '
     'generated through use of the ATLAS Platform by Al Thiqa personnel.'),
    ('"Research Data"',
     'means aggregate or anonymised Clinical Data used for peer-reviewed research, publication, or '
     'regulatory reporting purposes under Section 7 of this Agreement.'),
    ('"Authorised Users"',
     'means Al Thiqa personnel assigned one of the three access tiers defined in Section 3.2.'),
    ('"Personal Data" and "Processing"',
     'have the meanings given in UAE Federal Decree-Law No. 45 of 2021 (the PDPL).'),
]

for term, defn in definitions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(1.2)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_before      = Pt(3)
    p.paragraph_format.space_after       = Pt(3)
    run(p, term + '  ', bold=True, color=NAVY)
    run(p, defn)

# =============================================================================
# 2. PLATFORM ACCESS
# =============================================================================
section_heading(2, 'PLATFORM ACCESS - SCOPE AND TERM')

sub('2.1', 'Network Coverage.  ',
    'Adherence Inc. shall provision and maintain nine (9) independent ATLAS Workspaces, '
    'one per Al Thiqa location as listed in Schedule A, for the duration of this Agreement.')

sub('2.2', 'Instruments Licensed.  ',
    'All nine Workspaces shall have access to both the MMAS-8 and MAP instruments. '
    'The MMAS-8 is licensed under a separate instrument licence held by Donald E. Morisky, '
    'ScD, MSPH, PhD. Use of MMAS-8 under this Agreement is permitted solely within the '
    'ATLAS Platform and solely for Al Thiqa\'s clinical and research purposes as defined '
    'herein. Redistribution, extraction, or standalone use of the MMAS-8 outside the '
    'ATLAS Platform is not permitted.')

sub('2.3', 'Term.  ',
    'This Agreement commences on the Effective Date and continues for an initial term of '
    'twelve (12) months (the "Initial Term"), renewing automatically for successive '
    'twelve-month periods unless either Party provides written notice of non-renewal at '
    'least sixty (60) days before the end of the then-current term.')

sub('2.4', 'Onboarding.  ',
    'Adherence Inc. shall provide remote onboarding for up to two (2) Al Thiqa clinical '
    'leads within thirty (30) days of the Effective Date, covering platform setup, '
    'workspace configuration, Authorised User provisioning, and instrument administration protocols.')

# =============================================================================
# 3. AUTHORISED USERS
# =============================================================================
section_heading(3, 'AUTHORISED USERS AND ACCESS TIERS')

sub('3.1', 'Role Structure.  ',
    'Al Thiqa personnel shall be provisioned under three access tiers as set out in '
    'Section 3.2 below. No superadmin access is granted under this Agreement. '
    'Platform-level administration remains the sole responsibility of Adherence Inc.')

sub('3.2', 'Access Tiers.  ', '')

styled_table(
    ['Tier', 'Allocation', 'Permissions'],
    [
        ['Institution',
         'One (1) per network',
         'Full network-level visibility across all nine Al Thiqa locations. '
         'Intended for executive leadership (CEO, COO, CTO, Director level). '
         'Sees all child PI and clinician workspace data in aggregate and at '
         'individual location level. Network administration; user management '
         'across all sites; data export at network or per-location level.'],
        ['Principal Investigator (PI)',
         'Up to two (2) network-wide',
         'Full network-level visibility across all nine Al Thiqa locations, '
         'equivalent to Institution tier. Intended for clinical programme leads '
         'responsible for adherence operations. Sees all child clinician workspace '
         'data; research-grade data export; IRB-ready documentation; GAI '
         'contribution management; co-investigator status on research outputs '
         '(see Section 7). Clinician data flows upward to PI visibility automatically.'],
        ['Clinician',
         'Unlimited across the Network',
         'Assessment administration (MMAS-8 and MAP); patient record entry; '
         'individual score review within their assigned location. Assessment '
         'data is automatically visible to the PI and Institution tiers above. '
         'Clinicians do not have visibility into other locations or other '
         'clinicians\' records.'],
    ],
    col_widths=[1.4, 1.6, 3.5]
)

sub('3.3', 'Network Visibility Architecture.  ',
    'Both the Institution and PI tiers have full visibility across all nine '
    'Al Thiqa Network locations. Assessment data submitted by Clinician-tier '
    'users is automatically aggregated upward and visible to all PI and '
    'Institution users within the Network. This architecture enables executive '
    'oversight and programme leadership to monitor adherence performance '
    'across all sites without requiring separate logins per location.')

sub('3.4', 'User Management.  ',
    'Al Thiqa is responsible for promptly notifying Adherence Inc. of any change in '
    'personnel that requires access modification, suspension, or revocation. '
    'Adherence Inc. shall action such requests within two (2) business days.')

# =============================================================================
# 4. FEES
# =============================================================================
section_heading(4, 'FEES AND PAYMENT')

sub('4.1', 'Subscription Fees.  ',
    'Al Thiqa shall pay the following monthly subscription fees for each location, '
    'reflecting the clinical complexity and patient volume profile of each site type:')

styled_table(
    ['Location', 'Type', 'Monthly Fee (USD)', 'Annual Fee (USD)'],
    [
        ['[Specialty Pharmacy Name]', 'Specialty Pharmacy',    '299',             '3,588'],
        ['[Retail Locations 1-8] x8', 'Retail Pharmacy',      '149 per location', '1,788 per location'],
        ['NETWORK TOTAL',             '',                       '1,491 per month', '17,892 per year'],
    ],
    col_widths=[2.2, 1.5, 1.5, 1.5]
)

sub('4.2', 'Inclusions.  ',
    'The fees in Section 4.1 include all Workspace provisioning; MMAS-8 and MAP instrument '
    'access; all Authorised User tiers at the allocations specified in Section 3.2; UAE '
    'data residency infrastructure (Section 6); technical support (Section 5); and access '
    'to the research collaboration framework (Section 7).')

sub('4.3', 'Invoicing.  ',
    'Adherence Inc. shall issue invoices on the first business day of each calendar month. '
    'Payment is due within thirty (30) days of the invoice date.')

sub('4.4', 'Annual Pre-payment Discount.  ',
    'Al Thiqa may elect to pay the full annual fee in advance. Annual pre-payment attracts '
    'a ten percent (10%) discount, reducing the annual network fee to USD 16,103.')

sub('4.5', 'Fee Review.  ',
    'Fees shall remain fixed for the Initial Term. Any fee adjustment for renewal terms '
    'shall be notified in writing no later than ninety (90) days before the renewal date '
    'and shall not exceed five percent (5%) per annum.')

sub('4.6', 'Currency.  ',
    'All fees are denominated in United States Dollars. Al Thiqa may request invoicing in '
    'UAE Dirhams (AED) at the prevailing UAE Central Bank mid-rate on the invoice date.')

# =============================================================================
# 5. SUPPORT
# =============================================================================
section_heading(5, 'SUPPORT AND SERVICE LEVELS')

sub('5.1', 'Technical Support.  ',
    'Adherence Inc. shall provide email support at support@adherence.cc with a response '
    'commitment of one (1) business day for standard queries and four (4) hours for '
    'critical platform availability issues.')

sub('5.2', 'Availability.  ',
    'Adherence Inc. shall use commercially reasonable efforts to maintain ATLAS Platform '
    'availability of not less than 99.5% measured monthly, excluding scheduled maintenance '
    'windows communicated at least 48 hours in advance.')

sub('5.3', 'Training.  ',
    'Adherence Inc. shall provide up to two (2) additional remote training sessions per '
    'contract year at no additional charge, covering new features, research workflows, '
    'or clinical staff onboarding as requested by Al Thiqa.')

sub('5.4', 'Compliance Documentation.  ',
    'Adherence Inc. shall maintain and make available on request the ATLAS Data Security '
    'Fact Sheet, Data Processing Agreement, Technical and Organisational Measures '
    'attestation, and ADHICS v2.0 compliance documentation, available at '
    'atlas.adherence.cc/security.html.')

# =============================================================================
# 6. DATA PROTECTION
# =============================================================================
section_heading(6, 'DATA PROTECTION AND UAE DATA RESIDENCY')

sub('6.1', 'Data Residency.  ',
    'All Personal Data generated by Al Thiqa Authorised Users shall be processed and '
    'stored exclusively within the United Arab Emirates. Adherence Inc. confirms that '
    'Al Thiqa Workspace data is routed to and retained in AWS DynamoDB, me-central-1 '
    'region (Abu Dhabi, UAE) and is not replicated to or accessible from any non-UAE '
    'AWS infrastructure region.')

sub('6.2', 'PDPL Compliance.  ',
    'The Parties shall execute a Data Processing Agreement (DPA) as a separate Schedule '
    'to this Agreement prior to any Personal Data being submitted to the ATLAS Platform. '
    'The DPA governs all Processing of Personal Data and incorporates the sub-processor '
    'register confirming UAE data residency for all Al Thiqa workspace data.')

sub('6.3', 'ADHICS v2.0.  ',
    'Adherence Inc. confirms that the ATLAS Platform architecture is aligned with ADHICS '
    'v2.0 (Abu Dhabi Healthcare Information and Cyber Security Standard). Al Thiqa is '
    'responsible for ensuring its internal policies, staff training, and device security '
    'controls satisfy the organisational requirements of ADHICS v2.0 applicable to '
    'healthcare information processors.')

sub('6.4', 'Breach Notification.  ',
    'In the event of a confirmed or suspected Personal Data breach, Adherence Inc. shall '
    'notify Al Thiqa within seventy-two (72) hours of discovery in accordance with PDPL '
    'Article 14 and ADHICS v2.0 incident response controls IR-1 through IR-4.')

# =============================================================================
# 7. RESEARCH COLLABORATION
# =============================================================================
section_heading(7, 'RESEARCH COLLABORATION')

# Academic callout box
p_call = doc.add_paragraph()
p_call.paragraph_format.left_indent  = Cm(0.8)
p_call.paragraph_format.right_indent = Cm(0.8)
p_call.paragraph_format.space_before = Pt(8)
p_call.paragraph_format.space_after  = Pt(10)
shade_paragraph(p_call, BLUE_TINT)
run(p_call,
    'Research Rationale.  ',
    bold=True, color=NAVY)
run(p_call,
    'Al Thiqa\'s nine-location UAE pharmacy network provides access to a patient '
    'population that is substantially underrepresented in published medication adherence '
    'literature. The United Arab Emirates carries one of the highest chronic disease '
    'burdens globally, with diabetes prevalence estimated at approximately 17 percent '
    'and hypertension and dyslipidaemia widely prevalent across both national and '
    'expatriate populations. Adherence data generated through this Agreement, using the '
    'MMAS-8 and MAP instruments as primary endpoints, would constitute a material '
    'contribution to Gulf Arab adherence epidemiology. Al Thiqa\'s designated '
    'Principal Investigator-tier users are recognised as Co-Investigators on all '
    'resulting peer-reviewed publications, subject to satisfaction of ICMJE authorship criteria.',
    italic=True, color=NAVY)

sub('7.1', 'Research Partnership.  ',
    'The Parties recognise that the clinical dataset generated through Al Thiqa\'s use '
    'of the ATLAS Platform constitutes a significant research asset. The dataset covers '
    'medication adherence across a multi-site UAE pharmacy network with a '
    'demographically diverse patient population. The Parties agree to collaborate in '
    'good faith on peer-reviewed research utilising this dataset on the terms set out '
    'in this Section.')

sub('7.2', 'Priority Research Areas.  ',
    'The Parties have identified the following as priority research domains for initial collaboration:')

research_areas = [
    '(a)  Medication adherence patterns in UAE community pharmacy populations, '
         'stratified by chronic disease category including diabetes, hypertension, '
         'dyslipidaemia, and depression;',
    '(b)  Comparative adherence profiles across specialty versus retail pharmacy settings '
         'in the UAE, utilising the multi-site structure of the Al Thiqa Network;',
    '(c)  Validation and normative benchmarking of the MAP instrument in Gulf Arab '
         'patient populations, with reference to published MMAS-8 norms;',
    '(d)  Longitudinal assessment of the impact of pharmacist-led adherence counselling '
         'on MMAS-8 and MAP scores within the Al Thiqa Network.',
]
for area in research_areas:
    bullet(area)

sub('7.3', 'PI Seats and Research Access.  ',
    'The two (2) PI-tier Authorised Users designated by Al Thiqa under Section 3.2 shall '
    'have access to research-grade data export, IRB/IEC-ready documentation generated '
    'by the ATLAS Platform, and aggregate cross-network analytics. These users are '
    'recognised as Co-Investigators on any research outputs arising from Al Thiqa Network '
    'data, subject to meeting ICMJE authorship criteria.')

sub('7.4', 'Philip Morisky as Principal Investigator.  ',
    'Philip Morisky, MBA, as developer of the MMAS-8 and MAP instruments and Principal '
    'Investigator at Adherence Inc., shall serve as Principal Investigator or '
    'Co-Principal Investigator on any peer-reviewed publication utilising Al Thiqa '
    'Network data. Al Thiqa\'s designated PI-tier users shall be listed as '
    'Co-Investigators.')

sub('7.5', 'IRB/IEC Approvals.  ',
    'Any research using patient-level data requires prospective approval from the '
    'relevant UAE ethics committee, including the DOH Research Ethics Committee, '
    'Dubai Healthcare City Ethics Committee, or equivalent body. Al Thiqa\'s designated '
    'PI is responsible for initiating and maintaining the required ethics approval. '
    'Adherence Inc. shall provide all required platform validation documentation, '
    'including IQ/OQ/PQ, SRS, and RTM documents, to support the ethics application '
    'at no additional charge.')

sub('7.6', 'Data Ownership.  ',
    'Individual-level Clinical Data generated in Al Thiqa Workspaces remains the '
    'property of Al Thiqa. Adherence Inc. may use de-identified, aggregated Al Thiqa '
    'Network data as a contribution to the Global Adherence Index (GAI), an anonymised '
    'epidemiological dataset, subject to the condition that no Al Thiqa patient records '
    'are individually identifiable in any GAI output. Al Thiqa may opt out of GAI '
    'contribution by written notice at any time.')

sub('7.7', 'Publication Rights.  ',
    'Neither Party shall publish or present research based on Al Thiqa Network data '
    'without the prior written consent of the other Party, which shall not be '
    'unreasonably withheld. Draft manuscripts shall be shared with the other Party '
    'no less than thirty (30) days before submission. The reviewing Party shall '
    'provide written comments or approval within that period.')

sub('7.8', 'Funding.  ',
    'The Parties agree to explore research funding opportunities through the UAE '
    'Research Foundation, ADEK, the Abu Dhabi Department of Health, and pharmaceutical '
    'industry real-world evidence programmes. Any external funding obtained shall be '
    'managed under a separate grant agreement consistent with the funding body\'s requirements.')

# =============================================================================
# 8. INTELLECTUAL PROPERTY
# =============================================================================
section_heading(8, 'INTELLECTUAL PROPERTY')

sub('8.1', 'Platform IP.  ',
    'All intellectual property in the ATLAS Platform, including its software, '
    'instruments (subject to Section 8.2), scoring algorithms, and documentation, '
    'remains the exclusive property of Adherence Inc.')

sub('8.2', 'MMAS-8 Instrument.  ',
    'The MMAS-8 is the intellectual property of Donald E. Morisky, ScD, MSPH, PhD. '
    'Al Thiqa acknowledges that its right to administer the MMAS-8 under this Agreement '
    'is contingent on continued licensor authorisation and may not be sub-licensed '
    'or used independently of the ATLAS Platform.')

sub('8.3', 'Al Thiqa Data.  ',
    'All patient-level Clinical Data generated in Al Thiqa Workspaces remains the '
    'property of Al Thiqa. Adherence Inc. holds a limited licence to process that data '
    'solely for the purpose of operating the ATLAS Platform and, subject to Section 7.6, '
    'contributing to the GAI.')

sub('8.4', 'Research Outputs.  ',
    'Intellectual property in research outputs, including publications, datasets, and '
    'methodologies, arising from the collaboration under Section 7 shall be jointly '
    'owned by the Parties unless otherwise agreed in writing prior to the relevant '
    'research commencing.')

# =============================================================================
# 9. CONFIDENTIALITY
# =============================================================================
section_heading(9, 'CONFIDENTIALITY')

sub('9.1', '',
    'Each Party shall hold the other\'s Confidential Information in strict confidence '
    'and shall not disclose it to any third party without prior written consent, except '
    'to Representatives who have a need to know for the purposes of this Agreement and '
    'are bound by confidentiality obligations no less onerous than those herein.')

sub('9.2', '',
    'Confidential Information does not include information that: (a) is or becomes '
    'publicly available through no breach of this Agreement; (b) was already known to '
    'the Recipient before disclosure; (c) is independently developed without reference '
    'to the Discloser\'s Confidential Information; or (d) must be disclosed by law or '
    'regulatory order, provided the Discloser is given reasonable prior notice.')

sub('9.3', '',
    'Confidentiality obligations under this Section survive termination of this '
    'Agreement for a period of five (5) years.')

# =============================================================================
# 10. LIABILITY
# =============================================================================
section_heading(10, 'LIABILITY AND WARRANTIES')

sub('10.1', 'Platform Warranty.  ',
    'Adherence Inc. warrants that the ATLAS Platform will perform materially in '
    'accordance with its published specifications and that it will maintain the '
    'security controls described in the ATLAS Data Security Fact Sheet throughout '
    'the term of this Agreement.')

sub('10.2', 'Clinical Responsibility.  ',
    'The ATLAS Platform is a clinical decision-support tool. Al Thiqa\'s pharmacists '
    'and clinical staff retain full professional responsibility for all clinical '
    'decisions made in connection with patient care. Adherence Inc. does not provide '
    'clinical advice and the platform\'s output does not constitute a clinical recommendation.')

sub('10.3', 'Limitation of Liability.  ',
    'To the maximum extent permitted by applicable law, neither Party\'s aggregate '
    'liability to the other under or in connection with this Agreement shall exceed '
    'the total fees paid by Al Thiqa in the twelve (12) months immediately preceding '
    'the event giving rise to the claim.')

sub('10.4', 'Exclusions.  ',
    'Neither Party shall be liable for indirect, consequential, or punitive damages '
    'arising out of or in connection with this Agreement.')

# =============================================================================
# 11. TERMINATION
# =============================================================================
section_heading(11, 'TERMINATION')

sub('11.1', 'Termination for Convenience.  ',
    'Either Party may terminate this Agreement at the end of the Initial Term or any '
    'renewal term by providing sixty (60) days\' written notice before the relevant '
    'term end date.')

sub('11.2', 'Termination for Cause.  ',
    'Either Party may terminate immediately on written notice if the other Party: '
    '(a) commits a material breach that is not remedied within thirty (30) days of '
    'written notice; (b) becomes insolvent or enters administration; or (c) commits '
    'a breach of data protection obligations that cannot be remedied.')

sub('11.3', 'Effect of Termination.  ',
    'On termination: (a) all Workspace access shall be disabled within five (5) '
    'business days; (b) Adherence Inc. shall make Al Thiqa\'s Clinical Data available '
    'for export in machine-readable format for thirty (30) days following the '
    'termination date; (c) after that period, Adherence Inc. shall securely delete '
    'all Al Thiqa Personal Data from its systems and provide written confirmation '
    'of deletion; and (d) all accrued payment obligations survive termination.')

# =============================================================================
# 12. GOVERNING LAW
# =============================================================================
section_heading(12, 'GOVERNING LAW AND DISPUTE RESOLUTION')

sub('12.1', '',
    'This Agreement shall be governed by and construed in accordance with the laws '
    'of the United Arab Emirates and, to the extent applicable, the laws of '
    'the Emirate of Abu Dhabi.')

sub('12.2', '',
    'The Parties shall attempt to resolve any dispute arising out of or in connection '
    'with this Agreement through good-faith negotiation for a period of thirty (30) '
    'days following written notice of the dispute.')

sub('12.3', '',
    'If the dispute is not resolved within that period, it shall be referred to and '
    'finally resolved by arbitration under the rules of the Abu Dhabi Commercial '
    'Conciliation and Arbitration Centre (ADCCAC), with the seat of arbitration in '
    'Abu Dhabi and conducted in the English language.')

# =============================================================================
# 13. GENERAL
# =============================================================================
section_heading(13, 'GENERAL')

sub('13.1', 'Entire Agreement.  ',
    'This Agreement, together with its Schedules, constitutes the entire agreement '
    'between the Parties relating to its subject matter and supersedes all prior '
    'negotiations, representations, and agreements.')

sub('13.2', 'Amendments.  ',
    'No amendment to this Agreement is valid unless made in writing and signed by '
    'authorised representatives of both Parties.')

sub('13.3', 'Notices.  ',
    'Notices under this Agreement shall be in writing and delivered by email with '
    'read receipt or by courier to the addresses set out in Schedule B.')

sub('13.4', 'Waiver.  ',
    'Failure to enforce any provision of this Agreement does not constitute a waiver '
    'of the right to enforce it subsequently.')

sub('13.5', 'Severability.  ',
    'If any provision of this Agreement is held invalid or unenforceable, the '
    'remaining provisions continue in full force and effect.')

# =============================================================================
# SIGNATURE BLOCK
# =============================================================================
doc.add_page_break()
section_heading('', 'EXECUTION')

p_exec = doc.add_paragraph()
p_exec.paragraph_format.space_before = Pt(10)
p_exec.paragraph_format.space_after  = Pt(14)
run(p_exec, 'IN WITNESS WHEREOF, the Parties have executed this Agreement as of the Effective Date.')

sig_tbl = doc.add_table(rows=6, cols=2)
sig_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

sig_rows = [
    ['ADHERENCE CARTOGRAPHY, INC.',        'AL THIQA PHARMACY NETWORK'],
    ['',                                    ''],
    ['Signature: ___________________________', 'Signature: ___________________________'],
    ['Name:       Philip Morisky',          'Name:       ___________________________'],
    ['Title:      Founder & Principal Investigator', 'Title:      ___________________________'],
    ['Date:       ___________________________', 'Date:       ___________________________'],
]

for ri, (left, right) in enumerate(sig_rows):
    cells = sig_tbl.rows[ri].cells
    for cell, text in [(cells[0], left), (cells[1], right)]:
        p = cell.paragraphs[0]
        is_header = (ri == 0)
        run(p, text,
            bold=is_header, size=10 if is_header else 9.5,
            color=WHITE if is_header else MID)
        if is_header:
            shade_cell(cell, '0d1b3e')
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        set_cell_borders(cell)

spacer(16)

# =============================================================================
# SCHEDULES
# =============================================================================
add_rule('B8902E', 12, 0, 0)

def sched_title(title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run(p, title, bold=True, size=11, color=NAVY)
    return p

sched_title('SCHEDULE A - AL THIQA NETWORK LOCATIONS AND WORKSPACE ASSIGNMENTS')

styled_table(
    ['No.', 'Location Name', 'Type', 'Monthly Fee (USD)', 'Workspace ID'],
    [
        ['1', '[Specialty Pharmacy Name]', 'Specialty Pharmacy', '299',        'To be assigned'],
        ['2', '[Retail Location 1]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['3', '[Retail Location 2]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['4', '[Retail Location 3]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['5', '[Retail Location 4]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['6', '[Retail Location 5]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['7', '[Retail Location 6]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['8', '[Retail Location 7]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['9', '[Retail Location 8]',       'Retail Pharmacy',    '149',        'To be assigned'],
        ['',  'NETWORK TOTAL',             '',                   '1,491 / mo', ''],
    ],
    col_widths=[0.3, 2.0, 1.5, 1.5, 1.2]
)

sched_title('SCHEDULE B - NOTICES')

styled_table(
    ['Party', 'Contact', 'Address', 'Email'],
    [
        ['Adherence Cartography, Inc.', 'Philip Morisky',
         '100 Oceangate, 12th Floor, Long Beach, CA 90802, USA', 'info@adherence.cc'],
        ['Al Thiqa Pharmacy Network', '[Name / Title]',
         '[Address, Abu Dhabi, UAE]', '[Email]'],
    ],
    col_widths=[1.5, 1.4, 2.2, 1.4]
)

sched_title('SCHEDULE C - DATA PROCESSING AGREEMENT')
body('To be executed as a separate document prior to platform activation. '
     'Template available at atlas.adherence.cc/security.html. '
     'The DPA governs all Processing of Personal Data and incorporates the '
     'sub-processor register confirming UAE data residency (AWS DynamoDB, '
     'me-central-1, Abu Dhabi) for all Al Thiqa workspace data.')

# Footer
spacer(20)
add_rule('B8902E', 6, 0, 0)
add_rule('0d1b3e', 18, 0, 0)

p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(6)
run(p_foot,
    'ATLAS Clinical Platform  |  Adherence Cartography, Inc.  |  '
    'info@adherence.cc  |  atlas.adherence.cc  |  ATLAS v8.8.0  |  June 2026',
    size=7.5, color=DIM, font='Courier New')

p_disc = doc.add_paragraph()
p_disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_disc.paragraph_format.space_before = Pt(4)
run(p_disc,
    'This agreement should be reviewed by qualified UAE legal counsel before execution. '
    'MMAS-8 licensing terms (Section 8.2) should be confirmed with the instrument licensor prior to signature.',
    italic=True, size=7.5, color=DIM)

# =============================================================================
# SAVE
# =============================================================================
out = r'C:\Users\philm\documents\ATLAS_AlThiqa_Network_Agreement_v2.docx'
doc.save(out)
print(f'Saved: {out}')
